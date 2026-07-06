#!/usr/bin/env python3
"""Generate HemaGuide input docx files from supplementary case text."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document

PROJECT_ROOT = Path(__file__).resolve().parent.parent
SUPP_PATH = PROJECT_ROOT / "paper" / "hema_guide_supplementary.txt"
KB_DIR = PROJECT_ROOT / "kb_input" / "tumorboards"
QUERY_DIR = PROJECT_ROOT / "query_input"

CASE_START = re.compile(
    r"^(?P<id>(?:l_case_\d+|l_case_sim_\d+|ly_case_\d+|ly_case_sim_\d+|"
    r"mm_case_\d+|mm_case_sim_\d+))$"
)
CASE_END = re.compile(
    r"^(?:"
    r"[a-z_]+__ALL|"
    r"#{3,}|"
    r"={10,}\s*(?:gpt-|qwen|claude|gemini|deepseek|llama|mistral|o[0-9]-)"
    r")",
    re.IGNORECASE,
)

BENCHMARK_IDS = (
    {f"l_case_{i}" for i in range(1, 16)}
    | {f"ly_case_{i:02d}" for i in range(1, 16)}
    | {f"mm_case_{i}" for i in range(1, 16)}
)
SIM_IDS = {"l_case_sim_1", "ly_case_sim_1", "mm_case_sim_1"}


def parse_cases(text: str) -> dict[str, str]:
    cases: dict[str, str] = {}
    current_id: str | None = None
    current_lines: list[str] = []

    for line in text.splitlines():
        stripped = line.strip()
        start = CASE_START.match(stripped)
        if start:
            if current_id:
                cases[current_id] = "\n".join(current_lines).strip()
            current_id = start.group("id")
            current_lines = []
            continue

        if current_id is None:
            continue

        if CASE_END.match(stripped):
            cases[current_id] = "\n".join(current_lines).strip()
            current_id = None
            current_lines = []
            continue

        current_lines.append(line)

    if current_id:
        cases[current_id] = "\n".join(current_lines).strip()

    return cases


def write_docx(case_id: str, body: str, out_path: Path) -> None:
    doc = Document()
    doc.add_heading(f"Tumor Board Case — {case_id}", level=1)
    for paragraph in body.split("\n"):
        doc.add_paragraph(paragraph)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)


def main() -> None:
    text = SUPP_PATH.read_text(encoding="utf-8")
    cases = parse_cases(text)

    missing_benchmark = sorted(BENCHMARK_IDS - cases.keys())
    missing_sim = sorted(SIM_IDS - cases.keys())
    if missing_benchmark or missing_sim:
        raise SystemExit(
            f"Missing expected cases. benchmark={missing_benchmark}, sim={missing_sim}"
        )

    kb_written = 0
    query_written = 0

    for case_id in sorted(cases.keys()):
        body = cases[case_id]
        if not body:
            continue

        if case_id in BENCHMARK_IDS or case_id in SIM_IDS:
            write_docx(case_id, body, KB_DIR / f"{case_id}.docx")
            kb_written += 1

        if case_id in BENCHMARK_IDS:
            write_docx(case_id, body, QUERY_DIR / f"{case_id}.docx")
            query_written += 1

    print(f"Parsed {len(cases)} case blocks from supplementary text")
    print(f"Wrote {kb_written} docx files to {KB_DIR}")
    print(f"Wrote {query_written} docx files to {QUERY_DIR}")


if __name__ == "__main__":
    main()
