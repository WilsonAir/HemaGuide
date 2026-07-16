"""
Document extraction: two-step LLM-based extraction for clinical tumor board documents.

Step 1: Base section extraction (8 clinical sections, shared prompts)
        + prior_treatments derived from history via additional LLM call
Step 2: Molecular data extraction (3 fields, entity-specific prompts)

Architecture:
    Document → Entity Classification → Base Extraction → Prior Treatments → Molecular Extraction → Output

Key Design Choices:
    - Two-step extraction: Separates entity-agnostic base sections from entity-specific
      molecular data to enable config reuse across disease types
    - Document splitting: Optimizes context window by separating clinical narrative
      from molecular appendix before extraction
    - Caching: Avoids redundant LLM calls via file-based cache in extracted_data/
"""

import json
import logging
import re
import time
from pathlib import Path
from typing import Any
from datetime import datetime
import hashlib

from docx import Document as DocxDocument
from docx.table import Table
import yaml

from .llm import create_client, is_ollama_client
from .utils import extract_patient_id, log_result, supports_temperature
from .mol.cytogenetics import classify_fish_risk

logger = logging.getLogger(__name__)


# ============================================================================
# MODULE EXPORTS
# ============================================================================

__all__ = [
    'extract_document',
    'extract_document_cached',
    'extract_from_text',
    'extract_prior_treatments',
    'read_document',
    'read_docx',
    'build_extraction_prompt',
    'generate_document_id',
]


# ============================================================================
# CONSTANTS
# ============================================================================

DEFAULT_MODEL = "gpt-oss:120b"
DEFAULT_TEMPERATURE = 0.1
MAX_RETRIES = 2
DEBUG_DIR = Path('./extracted_data/json_debug')

DEFAULT_MOL_TB_PATTERNS = [
    r"(?:^|\n)\s*(Anhang[:\s–-]*)?Molekulargenetischer Befund\s*(?:\n|$)",
    r"(?:^|\n)\s*Molekulares Tumorboard\s*(?:\n|$)",
    r"(?:^|\n)\s*(Anhang[:\s–-]*)?Molekulardiagnostik\s*(?:\n|$)",
    r"(?:^|\n)\s*Molekular stratifizierte Therapieoptionen\s*(?:\n|$)"
]

ENTITY_CLASSIFICATION_SYSTEM = """Sie sind ein spezialisierter KI-Assistent für die Klassifizierung hämatologischer Erkrankungen.
Analysieren Sie den Dokumentkopf und identifizieren Sie die primäre Erkrankung."""

ENTITY_CLASSIFICATION_USER = """Basierend auf dem folgenden Dokumentkopf:

DOKUMENTKOPF:
{header_text}

Klassifizieren Sie diesen Fall in GENAU EINE der folgenden Entitäten:
{entity_list}

Antworten Sie NUR mit dem exakten Entitätsnamen, nichts anderes."""


# ============================================================================
# PUBLIC API
# ============================================================================

def extract_from_text(
    text: str,
    source_name: str = "pasted_case.txt",
    model: str = DEFAULT_MODEL,
    llm_mode: str | None = None,
    api_key: str | None = None,
) -> dict[str, Any]:
    """
    Extract structured data from raw case text (no file conversion).

    Args:
        text: Case vignette as plain text
        source_name: Logical source label used for document_id / patient_id
        model: LLM model to use
        llm_mode: LLM mode (openai, ollama-local, ollama-cloud, vllm)
        api_key: API key

    Returns:
        Dict with document_id, source_file, sections, enriched flag
    """
    cleaned = (text or "").strip()
    if not cleaned:
        raise ValueError("Case text is empty")

    source_name = Path(source_name).name or "pasted_case.txt"
    logger.info(f"Extraction {source_name}: from text ({len(cleaned)} chars)")

    clinical_text, molecular_text, is_mol_tb = split_document_at_molecular(cleaned)
    if is_mol_tb:
        logger.info(f"Extraction {source_name}: molTB detected, document split")
    else:
        logger.info(f"Extraction {source_name}: non-molTB, full document extraction")

    header = extract_header(clinical_text, max_lines=30)

    entities = load_entities()
    entity = classify_entity_from_header(
        header_text=header,
        entities=entities,
        model=model,
        llm_mode=llm_mode,
        api_key=api_key,
    )
    logger.info(f"Extraction {source_name}: entity={entity}")

    entity_slugs = load_entity_slugs()
    entity_slug = get_entity_slug(entity, entity_slugs)

    base_config_path = get_base_config_path()
    with open(base_config_path, 'r', encoding='utf-8') as f:
        base_config = yaml.safe_load(f)

    if is_mol_tb:
        user_prompt = build_extraction_prompt(base_config, clinical_text)
        sections = extract_sections(
            text=clinical_text,
            system_prompt=base_config['system_prompt'],
            user_prompt=user_prompt,
            json_schema=base_config['json_schema'],
            model=model,
            llm_mode=llm_mode,
            api_key=api_key,
        )
    else:
        user_prompt = build_extraction_prompt(base_config, cleaned)
        sections = extract_sections(
            text=cleaned,
            system_prompt=base_config['system_prompt'],
            user_prompt=user_prompt,
            json_schema=base_config['json_schema'],
            model=model,
            llm_mode=llm_mode,
            api_key=api_key,
        )

    history_text = sections.get('history', 'nicht vorhanden')
    if history_text and history_text.strip().lower() != 'nicht vorhanden':
        prior_treatments = extract_prior_treatments(
            history=history_text,
            model=model,
            llm_mode=llm_mode,
            api_key=api_key,
        )
    else:
        prior_treatments = []

    sections['prior_treatments'] = json.dumps(prior_treatments, ensure_ascii=False)

    if is_mol_tb:
        molecular_config_path = get_molecular_config_path(entity_slug)
        with open(molecular_config_path, 'r', encoding='utf-8') as f:
            molecular_config = yaml.safe_load(f)

        molecular = extract_molecular_data(
            molecular_text=molecular_text,
            config=molecular_config,
            model=model,
            llm_mode=llm_mode,
            api_key=api_key,
            entity_slug=entity_slug,
        )

        sections['mol_info'] = molecular['mol_info']
        sections['mol_fish'] = molecular['mol_fish']
        sections['mol_recom'] = molecular['mol_recom']
        sections['is_mol_tb'] = True
    else:
        sections['mol_info'] = 'nicht vorhanden'
        sections['mol_fish'] = 'nicht vorhanden'
        sections['mol_recom'] = 'nicht vorhanden'
        sections['is_mol_tb'] = False

    sections['entity'] = entity

    source_path = Path(source_name)
    output = {
        "document_id": generate_document_id(source_path),
        "source_file": source_path.name,
        "patient_id": extract_patient_id(source_path.name),
        "entity_slug": entity_slug,
        "extraction_timestamp": datetime.now().isoformat(),
        "sections": sections,
        "enriched": False,
    }

    log_result(logger, True, f"Extraction {source_path.name}", {
        "entity": entity_slug,
        "is_mol_tb": sections.get('is_mol_tb', False),
    })

    return output


def extract_document(
    file_path: Path,
    model: str = DEFAULT_MODEL,
    llm_mode: str | None = None,
    api_key: str | None = None
) -> dict[str, Any]:
    """
    Extract structured data from a document file (.docx or .txt).

    Args:
        file_path: Path to case file
        model: LLM model to use
        llm_mode: LLM mode (openai, ollama-local, ollama-cloud, vllm)
        api_key: API key

    Returns:
        Dict with document_id, source_file, sections, enriched flag
    """
    logger.info(f"Extraction {file_path.name}: reading document")
    text = read_document(file_path)
    return extract_from_text(
        text=text,
        source_name=file_path.name,
        model=model,
        llm_mode=llm_mode,
        api_key=api_key,
    )


def extract_document_cached(
    file_path: Path,
    cache_dir: Path,
    force_extract: bool = False,
    model: str = DEFAULT_MODEL,
    llm_mode: str | None = None,
    api_key: str | None = None
) -> dict[str, Any]:
    """Extract document with caching to avoid redundant LLM calls."""
    cache_dir.mkdir(parents=True, exist_ok=True)
    cache_path = cache_dir / f"{file_path.stem}.json"

    if not force_extract and cache_path.exists():
        with open(cache_path, 'r', encoding='utf-8') as f:
            return json.load(f)

    document = extract_document(
        file_path,
        model=model,
        llm_mode=llm_mode,
        api_key=api_key
    )

    with open(cache_path, 'w', encoding='utf-8') as f:
        json.dump(document, f, indent=2, ensure_ascii=False)

    return document


# ============================================================================
# EXTRACTION PIPELINE
# ============================================================================

def classify_entity_from_header(
    header_text: str,
    entities: list[str],
    model: str = DEFAULT_MODEL,
    llm_mode: str | None = None,
    api_key: str | None = None
) -> str:
    """Classify entity from document header. Runs before loading entity-specific config."""
    entity_list = ", ".join(entities)
    user_prompt = ENTITY_CLASSIFICATION_USER.format(
        header_text=header_text,
        entity_list=entity_list
    )

    messages = [
        {"role": "system", "content": ENTITY_CLASSIFICATION_SYSTEM},
        {"role": "user", "content": user_prompt}
    ]

    try:
        content = _call_llm(messages, model, llm_mode, api_key)

        # Exact match
        if content in entities:
            return content

        # Fuzzy match
        for entity in entities:
            if entity.lower() in content.lower() or content.lower() in entity.lower():
                return entity

        logger.warning(f"Entity '{content}' not in entities.txt — using fallback extraction.")
        return content.strip()

    except Exception as e:
        logger.warning(f"Entity classification failed ({e}). Using fallback.")
        return "Unbekannte Entität"


def extract_sections(
    text: str,
    system_prompt: str,
    user_prompt: str,
    json_schema: dict[str, Any],
    model: str = DEFAULT_MODEL,
    temperature: float = DEFAULT_TEMPERATURE,
    llm_mode: str | None = None,
    api_key: str | None = None
) -> dict[str, Any]:
    """Extract structured sections from text using LLM with JSON schema."""
    # For Ollama, inject schema example into prompt (no native schema support)
    if is_ollama_client(llm_mode):
        schema_example = _schema_to_example(json_schema)
        enhanced_prompt = (
            f"{user_prompt}\n\n"
            f"WICHTIG: Geben Sie die Antwort EXAKT in diesem JSON-Format zurück:\n"
            f"{schema_example}"
        )
        messages = [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": enhanced_prompt}
        ]
    else:
        messages = [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt}
        ]

    last_error = None
    for attempt in range(MAX_RETRIES + 1):
        content = _call_llm(
            messages, model, llm_mode, api_key,
            temperature=temperature,
            json_response=is_ollama_client(llm_mode),
            json_schema=json_schema if not is_ollama_client(llm_mode) else None
        )

        try:
            sections = json.loads(content)
        except json.JSONDecodeError:
            # Try stripping markdown code fences before giving up
            try:
                sections = json.loads(_clean_json_response(content))
            except json.JSONDecodeError:
                _save_debug_file(content, "malformed")
                if attempt < MAX_RETRIES:
                    last_error = f"JSON decode error on attempt {attempt + 1}"
                    continue
                raise

        is_valid, missing_keys = _validate_required_keys(sections, json_schema)
        if is_valid:
            return sections

        last_error = f"Missing required keys: {missing_keys}"
        if attempt < MAX_RETRIES:
            continue

        raise ValueError(
            f"Extraction failed after {MAX_RETRIES + 1} attempts. {last_error}. "
            f"Response was: {content[:500]}"
        )

    return sections


def extract_molecular_data(
    molecular_text: str,
    config: dict[str, Any],
    model: str = DEFAULT_MODEL,
    llm_mode: str | None = None,
    api_key: str | None = None,
    entity_slug: str = 'fallback'
) -> dict[str, str]:
    """Extract mol_info, mol_fish, mol_recom from analysis text."""
    default_result = {
        'mol_info': 'nicht vorhanden',
        'mol_fish': 'nicht vorhanden',
        'mol_recom': 'nicht vorhanden'
    }

    if not molecular_text or molecular_text.lower().strip() == 'nicht vorhanden':
        return default_result

    system_prompt = config.get('molecular_system_prompt', '')
    user_prompt_template = config.get('molecular_user_prompt', '')

    if not system_prompt or not user_prompt_template:
        return default_result

    user_prompt = user_prompt_template.format(analysis_data=molecular_text)
    messages = [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": user_prompt}
    ]

    try:
        content = _call_llm(
            messages, model, llm_mode, api_key,
            json_response=is_ollama_client(llm_mode)
        )
        cleaned = _clean_json_response(content)
        data = json.loads(cleaned)

        return _parse_molecular_fields(data, entity_slug)

    except json.JSONDecodeError:
        _save_debug_file(content, "molecular_malformed")
        return default_result
    except Exception:
        return default_result


def _parse_molecular_fields(data: dict, entity_slug: str) -> dict[str, str]:
    """Parse and validate molecular fields from LLM response."""
    result = {}

    # mol_info: list of variant dicts with 'gene' key
    mol_info = data.get('mol_info', 'nicht vorhanden')
    if isinstance(mol_info, list) and mol_info:
        valid = [v for v in mol_info if isinstance(v, dict) and v.get('gene')]
        result['mol_info'] = json.dumps(valid, ensure_ascii=False) if valid else 'nicht vorhanden'
    else:
        result['mol_info'] = 'nicht vorhanden'

    # mol_fish: list of FISH dicts with 'aberration' key
    mol_fish = data.get('mol_fish', 'nicht vorhanden')
    if isinstance(mol_fish, list) and mol_fish:
        valid = [f for f in mol_fish if isinstance(f, dict) and f.get('aberration')]
        if valid:
            valid = classify_fish_risk(valid, entity_slug=entity_slug)
            result['mol_fish'] = json.dumps(valid, ensure_ascii=False)
        else:
            result['mol_fish'] = 'nicht vorhanden'
    else:
        result['mol_fish'] = 'nicht vorhanden'

    # mol_recom: plain text string
    mol_recom = data.get('mol_recom', 'nicht vorhanden')
    if isinstance(mol_recom, str) and mol_recom.strip() and mol_recom.lower().strip() != 'nicht vorhanden':
        result['mol_recom'] = mol_recom.strip()
    else:
        result['mol_recom'] = 'nicht vorhanden'

    return result


def extract_prior_treatments(
    history: str,
    model: str = DEFAULT_MODEL,
    llm_mode: str | None = None,
    api_key: str | None = None
) -> list[str]:
    """
    Extract standardized therapy abbreviations from history text.

    Makes an additional LLM call to identify all therapies mentioned in the
    patient history and returns them in chronological order using standardized
    abbreviations from the hematology therapy reference list.

    Args:
        history: Extracted history text from base extraction
        model: LLM model to use
        llm_mode: LLM mode (openai, ollama-local, ollama-cloud)
        api_key: API key

    Returns:
        List of therapy abbreviation strings in chronological order
    """
    config_path = Path(__file__).parent.parent / "prompts" / "prior_treatments.yaml"
    with open(config_path, 'r', encoding='utf-8') as f:
        config = yaml.safe_load(f)

    system_prompt = config['system_prompt']
    user_prompt = config['user_prompt'].format(history=history)

    messages = [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": user_prompt}
    ]

    try:
        content = _call_llm(
            messages, model, llm_mode, api_key,
            json_response=is_ollama_client(llm_mode)
        )
        cleaned = _clean_json_response(content)
        result = json.loads(cleaned)

        if isinstance(result, list):
            return [t for t in result if isinstance(t, str) and t.strip()]

        if isinstance(result, dict):
            for v in result.values():
                if isinstance(v, list):
                    return [t for t in v if isinstance(t, str) and t.strip()]

        logger.warning(
            f"prior_treatments: unexpected response shape {type(result).__name__}, "
            f"content head: {content[:200]!r}"
        )
        return []

    except json.JSONDecodeError:
        logger.warning(f"prior_treatments: JSON parse failed, attempting line extraction")
        # Fallback: try to extract array from response
        match = re.search(r'\[.*?\]', content, re.DOTALL)
        if match:
            try:
                result = json.loads(match.group())
                if isinstance(result, list):
                    return [t for t in result if isinstance(t, str) and t.strip()]
            except json.JSONDecodeError:
                pass
        return []
    except Exception as e:
        logger.warning(f"prior_treatments extraction failed: {e}")
        return []


# ============================================================================
# DOCUMENT PREPROCESSING
# ============================================================================

def read_document(file_path: Path) -> str:
    """Read document text based on file extension (.docx or .txt)."""
    suffix = file_path.suffix.lower()
    if suffix == '.docx':
        return read_docx(file_path)
    if suffix == '.txt':
        return file_path.read_text(encoding='utf-8')
    raise ValueError(f"Unsupported file type: {suffix}. Supported: .docx, .txt")


def read_docx(file_path: Path) -> str:
    """Read text from .docx file preserving document order."""
    doc = DocxDocument(file_path)
    full_text = []

    para_map = {p._element: p for p in doc.paragraphs}
    table_map = {t._element: t for t in doc.tables}

    for element in doc.element.body:
        tag = element.tag.split('}')[-1] if '}' in element.tag else element.tag

        if tag == 'p':
            paragraph = para_map.get(element)
            if paragraph and paragraph.text.strip():
                full_text.append(paragraph.text)
        elif tag == 'tbl':
            table = table_map.get(element)
            if table:
                table_text = extract_table_text(table)
                if table_text:
                    full_text.extend(table_text)

    return '\n'.join(full_text)


def extract_table_text(table: Table) -> list[str]:
    """Extract text from docx table, preserving row structure."""
    text_lines = []

    for row in table.rows:
        row_texts = []
        processed_cells = set()

        for cell in row.cells:
            cell_key = id(cell._element)
            if cell_key in processed_cells:
                continue
            processed_cells.add(cell_key)

            if cell.text.strip():
                row_texts.append(cell.text)

        if row_texts:
            text_lines.append('\t'.join(row_texts))

    return text_lines


def split_document_at_molecular(
    text: str,
    patterns: list[str] | None = None
) -> tuple[str, str, bool]:
    """Split document at molecular header. Returns (clinical_text, molecular_text, is_mol_tb)."""
    patterns = patterns or DEFAULT_MOL_TB_PATTERNS

    for pattern in patterns:
        match = re.search(pattern, text, re.IGNORECASE | re.MULTILINE)
        if match:
            split_pos = match.start()
            if text[split_pos:split_pos+1] == '\n':
                split_pos += 1

            clinical_text = text[:split_pos].rstrip()
            molecular_text = text[split_pos:].lstrip()

            logger.debug(f"Document split at position {split_pos}, "
                        f"clinical={len(clinical_text)} chars, "
                        f"molecular={len(molecular_text)} chars")

            if len(clinical_text) < 100:
                logger.warning(f"Clinical part very short ({len(clinical_text)} chars) - "
                              f"molecular header may be at wrong position")

            return clinical_text, molecular_text, True

    return text, "", False


# ============================================================================
# CONFIGURATION & ROUTING
# ============================================================================

def load_entities(entities_file: Path | None = None) -> list[str]:
    """Load entity list from text file (one entity per line)."""
    if entities_file is None:
        entities_file = Path(__file__).parent.parent / "data" / "entities.txt"
    with open(entities_file, 'r', encoding='utf-8') as f:
        return [line.strip() for line in f if line.strip()]


def load_entity_slugs(slugs_file: Path | None = None) -> dict[str, str]:
    """Load entity name to YAML slug mapping."""
    if slugs_file is None:
        slugs_file = Path(__file__).parent.parent / "data" / "entity_slugs.json"
    with open(slugs_file, 'r', encoding='utf-8') as f:
        return json.load(f)


def get_entity_slug(entity: str, entity_slugs: dict[str, str]) -> str:
    """Get entity slug from entity name. Returns 'fallback' if not found."""
    slug = entity_slugs.get(entity)
    if not slug:
        logger.warning(f"No slug for entity '{entity}', using fallback extraction")
        return 'fallback'
    return slug


def get_base_config_path() -> Path:
    """Get path to shared base extraction config (extraction_base.yaml)."""
    config_path = Path(__file__).parent.parent / "prompts" / "extraction_base.yaml"
    if not config_path.exists():
        raise FileNotFoundError(f"Missing base extraction config: {config_path}")
    return config_path


def get_molecular_config_path(entity_slug: str) -> Path:
    """Get path to entity-specific molecular config with fallback."""
    config_path = Path(__file__).parent.parent / "prompts" / f"molecular_{entity_slug}.yaml"
    if config_path.exists():
        return config_path
    # Fallback to generic molecular config
    fallback_path = Path(__file__).parent.parent / "prompts" / "molecular_fallback.yaml"
    logger.info(f"Using fallback molecular config (no config for '{entity_slug}')")
    return fallback_path


def build_extraction_prompt(config: dict[str, Any], text: str) -> str:
    """Build extraction prompt from config template."""
    template = config.get('template', {})
    prompt_parts = [template.get('header', ''), ""]

    for i, section in enumerate(config['sections'], 1):
        prompt = config['section_prompts'].get(section, f"Extract {section}.")
        prompt_parts.append(f"{i}. {section}: {prompt}")

    section_keys = ", ".join(f'"{s}"' for s in config['sections'])
    for instruction in template.get('footer_instructions', []):
        if '{section_keys}' in instruction:
            instruction = instruction.format(section_keys=section_keys)
        prompt_parts.append(instruction)

    combined = "\n".join(prompt_parts)
    return f"{combined}\n\nText:\n{text}"


# ============================================================================
# INTERNAL HELPERS
# ============================================================================

def _call_llm(
    messages: list[dict],
    model: str,
    llm_mode: str | None,
    api_key: str | None,
    temperature: float = DEFAULT_TEMPERATURE,
    json_response: bool = False,
    json_schema: dict | None = None,
    max_retries: int = 5,
) -> str:
    """
    Make LLM call with unified Ollama/OpenAI interface and retry on rate limits.

    Args:
        messages: List of message dicts (system + user)
        model: Model identifier
        llm_mode: Backend mode ('openai', 'ollama-local', 'ollama-cloud')
        api_key: API key for the LLM service
        temperature: Sampling temperature
        json_response: Request JSON format (for Ollama)
        json_schema: JSON schema for structured output (for OpenAI)
        max_retries: Maximum retry attempts for rate limit errors

    Returns:
        Raw response content string
    """
    last_error = None

    for attempt in range(max_retries + 1):
        client = create_client(llm_mode, api_key)

        try:
            if is_ollama_client(llm_mode):
                params = {"model": model, "messages": messages}
                if json_response:
                    params["format"] = "json"
                if supports_temperature(model):
                    params["options"] = {"temperature": temperature}

                response = client.chat(**params)
                return response['message']['content'].strip()

            # OpenAI API
            params = {"model": model, "messages": messages}
            if json_schema:
                params["response_format"] = {"type": "json_schema", "json_schema": json_schema}
            if supports_temperature(model):
                params["temperature"] = temperature

            response = client.chat.completions.create(**params)
            return response.choices[0].message.content.strip()

        except Exception as e:
            error_str = str(e)
            # Check for rate limit error (429)
            if '429' in error_str or 'Too Many Requests' in error_str:
                last_error = e
                if attempt < max_retries:
                    wait_time = 2 ** attempt  # Exponential backoff: 1, 2, 4, 8, 16 seconds
                    logger.warning(f"Rate limit hit, retrying in {wait_time}s (attempt {attempt + 1}/{max_retries})")
                    time.sleep(wait_time)
                    continue
            raise

        finally:
            if hasattr(client, 'close'):
                client.close()

    # If we exhausted retries, raise the last error
    if last_error:
        raise last_error


def _save_debug_file(content: str, prefix: str = "malformed") -> Path:
    """Save malformed response to debug directory for inspection."""
    DEBUG_DIR.mkdir(parents=True, exist_ok=True)
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    debug_file = DEBUG_DIR / f"{prefix}_{timestamp}.txt"
    debug_file.write_text(content, encoding='utf-8')
    return debug_file


def _clean_json_response(content: str) -> str:
    """Remove markdown code block wrapper from JSON response."""
    cleaned = content.strip()
    if cleaned.startswith("```"):
        parts = cleaned.split("```")
        if len(parts) >= 2:
            cleaned = parts[1]
            if cleaned.startswith("json"):
                cleaned = cleaned[4:]
            cleaned = cleaned.strip()
    return cleaned


# ============================================================================
# UTILITY FUNCTIONS
# ============================================================================

def generate_document_id(file_path: Path) -> str:
    """Generate deterministic document ID from filename (SHA-256, 16 chars)."""
    filename = file_path.name
    hash_obj = hashlib.sha256(filename.encode('utf-8'))
    return hash_obj.hexdigest()[:16]


def extract_header(text: str, max_lines: int = 50) -> str:
    """Extract first N lines of document for entity classification."""
    lines = text.split('\n')[:max_lines]
    return '\n'.join(lines)


def _schema_to_example(json_schema: dict[str, Any]) -> str:
    """Convert JSON schema to example JSON string for prompt injection."""
    schema = json_schema.get('schema', {})
    properties = schema.get('properties', {})
    required_keys = schema.get('required', list(properties.keys()))

    example = {}
    for key in required_keys:
        prop = properties.get(key, {})
        prop_type = prop.get('type', 'string')
        if prop_type == 'array':
            example[key] = ["Beispiel1", "Beispiel2"]
        elif prop_type == 'string':
            if 'enum' in prop:
                example[key] = prop['enum'][0]
            else:
                example[key] = f"<{prop.get('description', key)}>"
        else:
            example[key] = f"<{key}>"

    return json.dumps(example, indent=2, ensure_ascii=False)


def _validate_required_keys(data: dict[str, Any], json_schema: dict[str, Any]) -> tuple[bool, list[str]]:
    """Validate response contains all required keys. Returns (is_valid, missing_keys)."""
    schema = json_schema.get('schema', {})
    required_keys = schema.get('required', [])
    missing = [k for k in required_keys if k not in data]
    return len(missing) == 0, missing
